from docx import Document 
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from datetime import datetime
import os
from app.config import Config
from docx.oxml import parse_xml
from PIL import Image

class CVGenerator:
    def __init__(self):
        self.template_path = Config.CV_TEMPLATE_PATH
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"No se encontró la plantilla en: {self.template_path}")
        
    def _agregar_imagen(self, doc, ruta_imagen, nombre):
        """Método para agregar imagen con formato correcto"""
        try:
            # Agregar párrafo para la imagen
            img_paragraph = doc.add_paragraph()
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Obtener dimensiones de la imagen
            with Image.open(ruta_imagen) as img:
                width, height = img.size
            
            # Calcular el ancho máximo (15 cm, manteniendo la proporción)
            max_width = Cm(15)
            aspect_ratio = height / width
            width = max_width
            height = width * aspect_ratio
            
            # Agregar la imagen centrada
            run = img_paragraph.add_run()
            run.add_picture(ruta_imagen, width=width)
            
            # Espacio después de la imagen
            doc.add_paragraph()
            
        except Exception as e:
            print(f"Error al agregar imagen {nombre}: {str(e)}")

    def generar(self, datos_personales, experiencias, output_path):
        try:
            # Usar la plantilla existente en lugar de crear un nuevo documento
            doc = Document(self.template_path)
            
            # DATOS PERSONALES
            titulo_dp = doc.add_paragraph('DATOS PERSONALES')
            titulo_dp.runs[0].font.bold = True
            titulo_dp.runs[0].font.size = Pt(16)
            
            # Datos personales con indentación
            campos = [
                ('APELLIDOS Y NOMBRES', datos_personales['nombres']),
                ('PROFESION', datos_personales['profesion']), 
                ('FECHA DE NACIMIENTO', datos_personales['fecha_nacimiento']),
                ('LUGAR DE NACIMIENTO', datos_personales['lugar_nacimiento']),
                ('REGISTRO C.I.P.', datos_personales.get('registro_cip', '')),
                ('RESIDENCIA', datos_personales.get('residencia', '')),
                ('TELÉFONO', datos_personales.get('telefono', '')),
                ('CORREO', datos_personales.get('correo', ''))
            ]
            
            # Tabla invisible para alineación
            table = doc.add_table(rows=len(campos), cols=3)
            table.allow_autofit = True
            
            # Ajustar ancho y quitar bordes
            for row in table.rows:
                for cell in row.cells:
                    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="auto" w:w="0"/>'))
                    cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'))
            
            # Agregar datos con indentación 
            for i, (label, valor) in enumerate(campos):
                cells = table.rows[i].cells
                
                # Primera celda: etiqueta
                label_para = cells[0].paragraphs[0]
                label_para.paragraph_format.left_indent = Inches(0.5)
                run = label_para.add_run(label)
                run.font.bold = True
                
                # Segunda celda: dos puntos centrados
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[1].paragraphs[0].add_run(' : ')
                
                # Tercera celda: valor
                cells[2].paragraphs[0].add_run(valor)

            doc.add_paragraph()  # Espacio entre secciones
            
            # EXPERIENCIA PROFESIONAL
            titulo_exp = doc.add_paragraph('EXPERIENCIA PROFESIONAL') 
            titulo_exp.runs[0].font.bold = True
            titulo_exp.runs[0].font.size = Pt(16)
            
            # Experiencias
            for exp in experiencias:
                campos_exp = [
                    ('Empresa', exp['empresa']),
                    ('Obra', exp['obra']),
                    ('Detalle de la obra', exp['detalle_obra']),
                    ('Monto del proyecto', exp['monto_proyecto']),
                    ('Cargo', exp['cargo']),
                    ('Periodo', f"{exp['fecha_inicio']} - {exp['fecha_fin']}"),
                    ('Propietario', exp['propietario'])
                ]
                
                # Tabla invisible para cada experiencia
                exp_table = doc.add_table(rows=len(campos_exp), cols=3)
                exp_table.allow_autofit = True
                
                # Quitar bordes
                for row in exp_table.rows:
                    for cell in row.cells:
                        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="auto" w:w="0"/>'))
                        cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'))
                
                # Agregar datos de experiencia
                for i, (label, valor) in enumerate(campos_exp):
                    cells = exp_table.rows[i].cells
                    
                    # Etiqueta con indentación
                    label_para = cells[0].paragraphs[0]
                    label_para.paragraph_format.left_indent = Inches(0.5)
                    run = label_para.add_run(label)
                    run.font.bold = True
                    
                    # Dos puntos centrados
                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cells[1].paragraphs[0].add_run(' : ')
                    
                    # Valor
                    cells[2].paragraphs[0].add_run(valor)
                
                # Funciones 
                if exp.get('funciones'):
                    func_para = doc.add_paragraph()
                    func_para.paragraph_format.left_indent = Inches(0.5)
                    func_para.add_run('Funciones:').bold = True
                    
                    funciones = exp['funciones']
                    if isinstance(funciones, str):
                        funciones = funciones.split(';')
                    
                    for funcion in funciones:
                        if funcion.strip():
                            bullet_para = doc.add_paragraph()
                            bullet_para.paragraph_format.left_indent = Inches(0.7)
                            bullet_para.paragraph_format.space_after = Pt(0)
                            bullet_para.paragraph_format.space_before = Pt(0)
                            check = bullet_para.add_run('✓')
                            check.font.size = Pt(9)
                            bullet_para.add_run(' ' + funcion.strip())

                # Imágenes de esta experiencia
                if 'documentos' in exp and exp['documentos']:
                    for img in exp['documentos']:
                        if os.path.exists(img['ruta']):
                            self._agregar_imagen(doc, img['ruta'], img['nombre'])
                
                doc.add_paragraph()  # Espacio entre experiencias

            # Documentos generales después de todas las experiencias
            if 'imagenes' in datos_personales and datos_personales['imagenes']:
                titulo_img = doc.add_paragraph('DOCUMENTOS ADJUNTOS')
                titulo_img.runs[0].font.bold = True 
                titulo_img.runs[0].font.size = Pt(16)
                
                for img in datos_personales['imagenes']:
                    if os.path.exists(img['ruta']):
                        self._agregar_imagen(doc, img['ruta'], img['nombre'])

            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error generando CV: {e}")
            return False